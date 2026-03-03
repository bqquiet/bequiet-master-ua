import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """Функція для встановлення кольору фону клітинки таблиці"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_resume():
    # Створюємо новий документ
    doc = docx.Document()
    
    # Встановлюємо вузькі поля
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Створюємо таблицю з 1 рядком та 2 колонками для макету
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Встановлюємо ширину колонок (ліва вужча, права ширша)
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(5.0)

    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]

    # Встановлюємо темно-сірий фон для лівої колонки
    set_cell_background(left_cell, "4A5568")

    # ================= ЛІВА КОЛОНКА =================
    
    # Місце для фото
    p_photo = left_cell.add_paragraph()
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_photo = p_photo.add_run("[ Вставте ваше фото сюди ]\n\n")
    run_photo.font.color.rgb = RGBColor(255, 255, 255)
    run_photo.font.size = Pt(12)

    # Особисті дані
    p_personal = left_cell.add_paragraph()
    run_personal_title = p_personal.add_run("Особисті дані\n")
    run_personal_title.bold = True
    run_personal_title.font.color.rgb = RGBColor(255, 255, 255)
    run_personal_title.font.size = Pt(14)
    
    # Лінія-розділювач
    p_line = left_cell.add_paragraph()
    r_line = p_line.add_run("______________________\n")
    r_line.font.color.rgb = RGBColor(255, 255, 255)

    personal_data = [
        ("Ім'я", "Андрій Бакалейко"),
        ("Адреса", "Тернопіль"),
        ("Номер телефону", "+380993279759"),
        ("Email", "dhdbcfdff@gmail.com")
    ]

    for key, val in personal_data:
        p = left_cell.add_paragraph()
        r_key = p.add_run(key + "\n")
        r_key.bold = True
        r_key.font.color.rgb = RGBColor(255, 255, 255)
        r_key.font.size = Pt(10)
        
        r_val = p.add_run(val)
        r_val.font.color.rgb = RGBColor(255, 255, 255)
        r_val.font.size = Pt(10)

    # Інтереси
    p_interests_title = left_cell.add_paragraph()
    run_interests_title = p_interests_title.add_run("\nІнтереси\n")
    run_interests_title.bold = True
    run_interests_title.font.color.rgb = RGBColor(255, 255, 255)
    run_interests_title.font.size = Pt(14)
    
    p_line2 = left_cell.add_paragraph()
    r_line2 = p_line2.add_run("______________________\n")
    r_line2.font.color.rgb = RGBColor(255, 255, 255)

    interests = [
        "Розробка ботів та автоматизація процесів",
        "Веб-розробка (Frontend + Backend)",
        "Створення внутрішніх інструментів та скриптів для бізнесу",
        "Ігрові ком'юніті та RP-проєкти"
    ]

    for interest in interests:
        p = left_cell.add_paragraph(style='List Bullet')
        r = p.add_run(interest)
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)

    # ================= ПРАВА КОЛОНКА =================
    
    # Ім'я
    p_name = right_cell.add_paragraph()
    run_name = p_name.add_run("Андрій Бакалейко")
    run_name.bold = True
    run_name.font.size = Pt(24)

    # Профіль
    p_profile = right_cell.add_paragraph()
    r_prof1 = p_profile.add_run("Андрій — розробник програмного забезпечення.\n")
    r_prof1.italic = True
    r_prof1.bold = True
    
    r_prof2 = p_profile.add_run(
        "Я студент спеціальності «Інженерія програмного забезпечення». "
        "Паралельно з навчанням працюю над реальними задачами та власними проєктами. "
        "Для мене головне — зробити робочий і зрозумілий продукт, а не просто написати код.\n"
        "Перед початком роботи я детально розбираюсь у задачі, продумую як краще її реалізувати "
        "та пропоную варіанти, якщо бачу можливість зробити систему простішою або ефективнішою. "
        "Пишу код акуратно, логічно та без зайвої складності."
    )
    r_prof2.font.size = Pt(10)

    # Досвід роботи
    p_exp_title = right_cell.add_paragraph()
    r_exp_title = p_exp_title.add_run("\nДосвід роботи")
    r_exp_title.bold = True
    r_exp_title.font.size = Pt(16)
    
    p_line3 = right_cell.add_paragraph()
    p_line3.add_run("____________________________________________________________")

    # Веб-розробка
    p_web = right_cell.add_paragraph()
    r_web_title = p_web.add_run("Веб-розробка (Frontend + Backend)\n")
    r_web_title.bold = True
    r_web_desc = p_web.add_run("Створюю сайти та веб-застосунки з нуля.\n")
    r_web_desc.font.size = Pt(10)
    
    r_front = p_web.add_run("Frontend:\n")
    r_front.bold = True
    r_front.font.size = Pt(10)
    
    front_items = ["адаптивна верстка", "створення зручного інтерфейсу", "реалізація логіки взаємодії", "підключення до серверної частини"]
    for item in front_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    p_back = right_cell.add_paragraph()
    r_back = p_back.add_run("Backend:\n")
    r_back.bold = True
    r_back.font.size = Pt(10)
    
    back_items = ["обробка даних", "створення API", "робота з базами даних", "авторизація користувачів", "інтеграція сторонніх сервісів"]
    for item in back_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    p_web_summary = right_cell.add_paragraph()
    p_web_summary.add_run("Можу зробити як простий лендінг, так і повноцінний веб-проєкт із логікою та системою керування даними.").font.size = Pt(10)

    # Боти
    p_bots = right_cell.add_paragraph()
    r_bots_title = p_bots.add_run("\nTelegram та Discord боти (Python)\n")
    r_bots_title.bold = True
    
    bot_items = ["автоматизація процесів", "системи модерації", "ролі та рівні доступу", "тікет-системи", "інтеграція з API", "сповіщення та моніторинг", "кастомна логіка під сервер або бізнес"]
    for item in bot_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    # Автоматизація
    p_auto = right_cell.add_paragraph()
    r_auto_title = p_auto.add_run("\nАвтоматизація та скрипти (Python)\n")
    r_auto_title.bold = True
    
    auto_items = ["автоматизація рутинних задач", "обробка тексту та файлів", "парсинг даних", "написання скриптів під конкретні потреби", "з’єднання різних сервісів між собою"]
    for item in auto_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    # C++
    p_cpp = right_cell.add_paragraph()
    r_cpp_title = p_cpp.add_run("\nПрограми на C++\n")
    r_cpp_title.bold = True
    
    cpp_items = ["створення консольних програм", "реалізація алгоритмів", "логічні та прикладні задачі", "структурована організація коду"]
    for item in cpp_items:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(10)

    # Освіта
    p_edu_title = right_cell.add_paragraph()
    r_edu_title = p_edu_title.add_run("\nОсвіта і кваліфікації")
    r_edu_title.bold = True
    r_edu_title.font.size = Pt(16)
    
    p_line4 = right_cell.add_paragraph()
    p_line4.add_run("____________________________________________________________")

    p_edu = right_cell.add_paragraph()
    p_edu.add_run("Студент першого курсу спеціальності «Інженерія програмного забезпечення».\nПаралельно працюю над практичними проєктами і вдосконалюю навички програмування.").font.size = Pt(10)

    # Навички
    p_skills_title = right_cell.add_paragraph()
    r_skills_title = p_skills_title.add_run("\nНавички")
    r_skills_title.bold = True
    r_skills_title.font.size = Pt(16)
    
    p_line5 = right_cell.add_paragraph()
    p_line5.add_run("____________________________________________________________")

    skills = ["Python (боти, автоматизація, скрипти)", "C++ (програми, алгоритми)", "Веб-розробка (Frontend + Backend)", "Git / GitHub", "Робота з API", "Дебагінг та тестування", "Організація проєктів та чистий код"]
    for skill in skills:
        p = right_cell.add_paragraph(style='List Bullet')
        p.add_run(skill).font.size = Pt(10)

    # Рекомендації
    p_rec_title = right_cell.add_paragraph()
    r_rec_title = p_rec_title.add_run("\nРекомендації")
    r_rec_title.bold = True
    r_rec_title.font.size = Pt(16)
    
    p_line6 = right_cell.add_paragraph()
    p_line6.add_run("____________________________________________________________")

    p_rec = right_cell.add_paragraph()
    p_rec.add_run("Готовий обговорювати ваш проєкт і допомогти реалізувати його з нуля.\nЯкщо є ідея — можу запропонувати оптимальне технічне рішення і план реалізації.").font.size = Pt(10)

    # Збереження
    filename = "resume_andriy.docx"
    doc.save(filename)
    print(f"Резюме успішно збережено у файл '{filename}'")

if __name__ == "__main__":
    create_resume()
